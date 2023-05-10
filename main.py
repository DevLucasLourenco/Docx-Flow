import docx
import os
    
    
    
class AdaptacaoWord():
        
   
    def save_file(self, nome_do_arquivo_saida):
        doc.save(caminho + sinal[0] + nome_do_arquivo_saida)
        
        self.nome_arquivo = nome_do_arquivo_saida
        
        self._mostrar_nome()


    def find_archive(self,nome):
        
        global doc
        global caminho
        global sinal
        
        
        caminho = os.getcwd()
        sinal = '\\'
        nome_arquivo = sinal[0] + nome
        diretorio = caminho + nome_arquivo
        doc = docx.Document(diretorio) 
        
            
    def edit_file(self, lista_texto_para_mudar, lista_texto_substituido):
        
        for i, item in enumerate(lista_texto_para_mudar):
            
            for paragrafo in doc.paragraphs:
                if item in paragrafo.text:
                    # identifica o texto que será alterado e separa em três partes. O anterior à parte que será separada, o posterior à ela e a própria.
                    texto_anterior, texto_posterior = paragrafo.text.split(item)
                    
                    # limpa tudo e adiciona o texto formatado e alterado, unindo os trechos os quais haviam sido separados.
                    paragrafo.clear()
                    paragrafo.add_run(texto_anterior)
                    trecho = paragrafo.add_run(lista_texto_substituido[i])
                    trecho.bold = True
                    paragrafo.add_run(texto_posterior)
                    
                    
    def _mostrar_nome(self):
        print(f'O arquivo foi salvo com o nome de {self.nome_arquivo}')
        


## EXECUÇÃO     
lista_procurada = ['Valor1','Valor2','Valor3']
lista_substituicao = ['VALOR ALTERADO1','VALOR ALTERADO2','VALOR ALTERADO3']

adaptar = AdaptacaoWord()

adaptar.find_archive('teste.docx')
adaptar.edit_file(lista_procurada, lista_substituicao)
adaptar.save_file(nome_do_arquivo_saida='TesteTratado.docx')